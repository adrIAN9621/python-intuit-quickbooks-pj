from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()
# profile picture
document.add_picture('WhatsApp_Image_2021-08-30_at_04.31.21.jpeg', width=Inches(2.0))

# name phone number and email
name = input('what is your name? ')
speak('Hello' + name + 'how are you today')
speak('what is your phone number? ')
phone_number = input('what is your phone number? ')
speak('your phone number is' + phone_number)
speak('what is your email? ')
email = input('what is your email? ')
speak('your email is' + email)
document.add_paragraph(name + ' | ' + phone_number + ' | ' + email)
# about me
document.add_heading('About me')
document.add_paragraph(input('Tell about yourself  '))

# work experiance
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To Date')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' +  to_date + '\n').italic = True

experience_detalis = input('Describe your experience at ' + company)
p.add_run(experience_detalis)
# more experiences
while True:
    has_more_experience = input('Do you are more experience? Yer or No ')
    if has_more_experience.lower() == 'yes ':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' +  to_date + '\n').italic = True

        experience_detalis = input('Describe your experience at ' + company + '  ' )
        p.add_run(experience_detalis)
    else:
         break

# Skil list
document.add_heading('Skil list')
skil_list = input('your skil is? ')
p = document.add_paragraph(skil_list)
p.style = 'List Bullet'


while True:
    has_more_skils = input('Do you are more skils? Yes or No')
    if has_more_skils.lower() == 'yes':
        skil_list = input('your skil is? ')
        p = document.add_paragraph(skil_list)
        p.style = 'List Bullet'
    else:
        break

#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generator usung Adis co de"

document.save('cv.docx')